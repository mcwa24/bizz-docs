import os
import logging
from pathlib import Path
from typing import Optional

# PDF parsing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PyPDF2 not available. PDF parsing will be disabled.")

# Word parsing
try:
    from docx import Document
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    logging.warning("python-docx not available. Word parsing will be disabled.")

logger = logging.getLogger(__name__)

class DocumentParser:
    """
    Parser for various document formats: PDF, Word, TXT, Markdown
    """
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._parse_pdf if PDF_AVAILABLE else None,
            '.docx': self._parse_word if WORD_AVAILABLE else None,
            '.txt': self._parse_text,
            '.md': self._parse_text,
            '.markdown': self._parse_text
        }
    
    def parse_document(self, file_path: str) -> str:
        """
        Parse a document and return its text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        parser_func = self.supported_formats[file_extension]
        if parser_func is None:
            raise ValueError(f"Parser not available for format: {file_extension}")
        
        try:
            logger.info(f"Parsing document: {file_path}")
            text_content = parser_func(file_path)
            logger.info(f"Successfully parsed {file_path}, extracted {len(text_content)} characters")
            return text_content
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF document using PyPDF2"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is not available")
        
        text_content = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                
                return '\n\n'.join(text_content)
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise
    
    def _parse_word(self, file_path: Path) -> str:
        """Parse Word document using python-docx"""
        if not WORD_AVAILABLE:
            raise ImportError("python-docx is not available")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n\n'.join(text_content)
        except Exception as e:
            logger.error(f"Error parsing Word document {file_path}: {str(e)}")
            raise
    
    def _parse_text(self, file_path: Path) -> str:
        """Parse plain text files (TXT, MD, Markdown)"""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error parsing text file {file_path}: {str(e)}")
                raise
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return [ext for ext, parser in self.supported_formats.items() if parser is not None]
    
    def is_format_supported(self, file_path: str) -> bool:
        """Check if a file format is supported"""
        file_extension = Path(file_path).suffix.lower()
        return file_extension in self.supported_formats and self.supported_formats[file_extension] is not None
